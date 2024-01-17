from hashlib import sha256
import json
import time
from PIL import Image,ImageTk
from flask import Flask
import tkinter as tk
import re
from tkinter import filedialog,messagebox,ttk
from docx import Document
from docx.shared import Pt, RGBColor
import os

class LSB():
    def encode_image(self,img, msg):
        length = len(msg)
        if length > 255:
            print("text too long! (don't exeed 255 characters)")
            return False
        encoded = img.copy()
        width, height = img.size
        index = 0
        for row in range(height):
            for col in range(width):
                if img.mode != 'RGB':
                    r, g, b ,a = img.getpixel((col, row))
                elif img.mode == 'RGB':
                    r, g, b = img.getpixel((col, row))
                if row == 0 and col == 0 and index < length:
                    asc = length
                elif index <= length:
                    c = msg[index -1]
                    asc = ord(c)
                else:
                    asc = b
                encoded.putpixel((col, row), (r, g , asc))
                index += 1
        return encoded

    def decode_image(self,img):
        width, height = img.size
        msg = ""
        index = 0
        for row in range(height):
            for col in range(width):
                if img.mode != 'RGB':
                    r, g, b ,a = img.getpixel((col, row))
                elif img.mode == 'RGB':
                    r, g, b = img.getpixel((col, row))
                if row == 0 and col == 0:
                    length = b
                elif index <= length:
                    msg += chr(b)
                index += 1
        return msg
class LeastSB:
    def d_image(self, img):
        width, height = img.size
        msg = ""
        index = 0
        for row in range(height):
            for col in range(width):
                if img.mode != 'RGB':
                    r, g, b, a = img.getpixel((col, row))
                elif img.mode == 'RGB':
                    r, g, b = img.getpixel((col, row))
                if row == 0 and col == 0:
                    length = b
                elif index <= length:
                    msg += chr(b)
                index += 1
        return msg

def check_image_for_fingerprint():
    def sel_image():
        nonlocal entry
        file_path = filedialog.askopenfilename(filetypes=[("Image Files", "*.png;*.jpg;*.jpeg;*.gif;*.bmp")])
        entry.delete(0, tk.END)
        entry.insert(0, file_path)

    def check_image():
        image_path = entry.get()
        try:
            image_to_check = Image.open(image_path)
            lsb_decoder = LeastSB()
            hidden_message = lsb_decoder.d_image(image_to_check)

            valid_fingerprint_format = re.match(r"^[a-f0-9]{64}$", hidden_message)

            if valid_fingerprint_format:
                result_label.config(text="The image contains a valid digital fingerprint.\nFingerprint: " + hidden_message)
            else:
                result_label.config(text="The image does not contain a valid digital fingerprint.")
        except Exception as e:
            result_label.config(text="Error: " + str(e))

    fingerprint_check_window = tk.Toplevel(root)
    fingerprint_check_window.title("Image Fingerprint Checker")

    label = tk.Label(fingerprint_check_window, text="Select an image:")
    label.pack()

    entry = tk.Entry(fingerprint_check_window, width=50)
    entry.pack()

    browse_button = tk.Button(fingerprint_check_window, text="Browse", command=sel_image)
    browse_button.pack()

    check_button = tk.Button(fingerprint_check_window, text="Check Image", command=check_image)
    check_button.pack()

    result_label = tk.Label(fingerprint_check_window, text="")
    result_label.pack()
class Block:
    def __init__(self, ID, images, timestamp, previousHash):
        self.ID = ID
        self.images = images
        self.timestamp = timestamp
        self.previousHash = previousHash
        self.nonce = 0

    def compute_hash(self):
        data = json.dumps(self.__dict__, sort_keys=True)
        return sha256(data.encode()).hexdigest()


class Blockchain:
    difficulty = 4

    def __init__(self):
        self.chain = []
        self.create_firstBlock()

    def create_firstBlock(self):
        firstBlock = Block(0, [], time.time(), "0")
        firstBlock.hash = firstBlock.compute_hash()
        self.chain.append(firstBlock)

    @property
    def lastBlock(self):
        return self.chain[-1]

    def add_block(self, block, proof):

        previous = self.lastBlock.hash

        if previous != block.previousHash:
            return False

        if not self.is_valid(block, proof):
            return False

        block.hash = proof
        self.chain.append(block)
        return True

    def is_valid(self, block, hashN):

        return (hashN.startswith('0' * Blockchain.difficulty) and
                hashN == block.compute_hash())

    def pOw(self, block):

        block.nonce = 0

        computedHash = block.compute_hash()
        while not computedHash.startswith('0' * Blockchain.difficulty):
            block.nonce += 1
            computedHash = block.compute_hash()

        return computedHash

    def add_new_transaction(self, image,seller,buyer):
        self.unconfirmedImages = image
        self.seller = seller
        self.buyer = buyer

    def mineBlock(self):
        imageS = Image.open(self.unconfirmedImages)
        fingerPrint = LSB().decode_image(imageS)
        if not self.unconfirmedImages:
            return False
        if fingerPrint != self.seller.fingerprint:
            print("false")
            return False
        imageB = LSB().encode_image(image2, self.buyer.fingerprint)
        imageB.save("encoded.png")
        lastBlock = self.lastBlock

        newBlock = Block(ID=lastBlock.ID + 1,
                         images=self.unconfirmedImages,
                         timestamp=time.time(),
                         previousHash=lastBlock.hash)
        proof = self.pOw(newBlock)
        self.add_block(newBlock, proof)

        self.unconfirmedImages = []
        return newBlock.ID

class User:
    def __init__(self,name,fingerprint):
        self.name = name
        self.fingerprint = sha256(fingerprint.encode()).hexdigest()

def show_message(message):
    messagebox.showinfo("Message", message)
def select_image():
    file_path = filedialog.askopenfilename()
    if file_path:
        global image2
        image2 = Image.open(file_path)
        status_label.config(text="Image selected")
        photo = ImageTk.PhotoImage(image2)
        img_label.config(image=photo)
        img_label.image2 = photo
        show_message("Image loaded successfully")

def process_image():
    def submit_user_data():
        seller_name = seller_name_entry.get()
        seller_fingerprint = seller_fingerprint_entry.get()
        buyer_name = buyer_name_entry.get()
        buyer_fingerprint = buyer_fingerprint_entry.get()

        seller1 = User(seller_name, seller_fingerprint)
        buyer1 = User(buyer_name, buyer_fingerprint)

        if image2:
            image1 = LSB().encode_image(image2, seller1.fingerprint)
            image1.save("encoded.png")
            blockchain.add_new_transaction("encoded.png", seller1, buyer1)
            blockchain.mineBlock()
            show_message("Image Processed, Please Check the root folder")
        else:
            show_message("Please select an image first.")

    # Tkinter UI for user input
    user_input_window = tk.Toplevel(root)
    user_input_window.title("Enter User Data")

    seller_name_label = tk.Label(user_input_window, text="Admin(Seller) Name:")
    seller_name_label.pack()
    seller_name_entry = tk.Entry(user_input_window)
    seller_name_entry.pack()

    seller_fingerprint_label = tk.Label(user_input_window, text="Admin(Seller) Fingerprint:")
    seller_fingerprint_label.pack()
    seller_fingerprint_entry = tk.Entry(user_input_window)
    seller_fingerprint_entry.pack()

    buyer_name_label = tk.Label(user_input_window, text="User(Buyer) Name:")
    buyer_name_label.pack()
    buyer_name_entry = tk.Entry(user_input_window)
    buyer_name_entry.pack()

    buyer_fingerprint_label = tk.Label(user_input_window, text="User(Buyer) Fingerprint:")
    buyer_fingerprint_label.pack()
    buyer_fingerprint_entry = tk.Entry(user_input_window)
    buyer_fingerprint_entry.pack()

    submit_button = tk.Button(user_input_window, text="Submit", command=submit_user_data)
    submit_button.pack()

blockchain = Blockchain()
image2 = None  # Initialize image2 as None

# Tkinter UI setup
root = tk.Tk()
root.title("Digi-llect protector")

# Styling using the 'clam' theme
style = ttk.Style()
style.theme_use("clam")

# Creating a container frame for buttons
button_frame = ttk.Frame(root, padding=10)
button_frame.pack()

# Adding icons for buttons
select_icon = Image.open('select_icon.png').resize((25, 25))
select_photo = ImageTk.PhotoImage(select_icon)

process_icon = Image.open('process_icon.png').resize((25, 25))
process_photo = ImageTk.PhotoImage(process_icon)

fingerprint_icon = Image.open('fingerprint_icon.png').resize((25, 25))
fingerprint_photo = ImageTk.PhotoImage(fingerprint_icon)

# Select Image button
select_button = ttk.Button(button_frame, text="Select Image", command=select_image, image=select_photo, compound=tk.LEFT)
select_button.image = select_photo  # Retain reference to the image
select_button.pack(side=tk.LEFT, padx=10)

# Process Image button
process_button = ttk.Button(button_frame, text="Process Image", command=process_image, image=process_photo, compound=tk.LEFT)
process_button.image = process_photo  # Retain reference to the image
process_button.pack(side=tk.LEFT, padx=10)

# Check Fingerprint button
third_button = ttk.Button(button_frame, text="Check Fingerprint", command=check_image_for_fingerprint, image=fingerprint_photo, compound=tk.LEFT)
third_button.image = fingerprint_photo  # Retain reference to the image
third_button.pack(side=tk.LEFT, padx=10)

image_frame = tk.Frame(root, bd=2, relief=tk.SUNKEN)
image_frame.pack(padx=20, pady=(0, 10))

photo = None
img_label = tk.Label(image_frame)
img_label.pack()

status_label = tk.Label(root, text="No image loaded", bd=1, relief=tk.SUNKEN, anchor=tk.W)
status_label.pack(fill=tk.X, padx=20)
def close_app():
    if root:
        root.quit()  # Stop the mainloop
        
root.mainloop()
result = messagebox.askyesno("Continue with Doc-Ellect(BETA)", "Continue with Doc-Ellect, Do you want to continue with Doc-Ellect Protector?(BETA)")
if result:
    # If the user chooses to continue, execute the following code
    class Tdoc:
        def __init__(self, root):
            self.root = root
            self.root.title("Doc-Ellect Protector")

            # ... (rest of the Tdoc class remains the same) ...
            self.frame = tk.Frame(self.root)
            self.frame.pack(padx=20, pady=20)

            self.select_button = tk.Button(self.frame, text="Select File", command=self.select_file)
            self.select_button.grid(row=0, column=0, pady=5)

            self.entry = tk.Entry(self.frame, width=50)
            self.entry.grid(row=0, column=1, padx=5, pady=5)

            self.watermark_label = tk.Label(self.frame, text="Enter Watermark Text:")
            self.watermark_label.grid(row=1, column=0)

            self.watermark_entry = tk.Entry(self.frame, width=50)
            self.watermark_entry.grid(row=1, column=1, padx=5, pady=5)

            self.add_button = tk.Button(self.frame, text="Add Watermark", command=self.add_watermark)
            self.add_button.grid(row=2, columnspan=2, pady=10)

            self.status_label = tk.Label(self.frame, text="", fg="green")
            self.status_label.grid(row=3, columnspan=2)

        def select_file(self):
            filepath = filedialog.askopenfilename()
            self.entry.delete(0, tk.END)
            self.entry.insert(0, filepath)

        def add_watermark(self):
            filepath = self.entry.get()
            watermark_text = self.watermark_entry.get()

            if not filepath or not watermark_text:
                self.status_label.config(text="Please select a file and enter watermark text.")
                return

            output_file = os.path.splitext(filepath)[0] + "_watermarked.docx"

            doc = Document(filepath)

            for section in doc.sections:
                header = section.header
                paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                run = paragraph.add_run()
                run.text = watermark_text
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(128, 128, 128)

            doc.save(output_file)
            self.status_label.config(text=f"Watermark added. Saved as {output_file}")
            pass

    def main():
        root = tk.Tk()
        app = Tdoc(root)
        root.mainloop()

        if __name__ == "__main__":
            main()
else:
    close_app()

def main():
    root = tk.Tk()
    app = Tdoc(root)
    root.mainloop()

if __name__ == "__main__":
    main()



app = Flask(__name__)
@app.route('/chain', methods=['GET'])
def get_chain():
    chain_data = []
    for block in blockchain.chain:
        chain_data.append(block.__dict__)
    return json.dumps({"length": len(chain_data),
                       "chain": chain_data})

app.run(debug=True, port=3000)


