import tkinter as tk
from tkinter import filedialog
from docx import Document

def select_file():
    filepath = filedialog.askopenfilename()
    entry.delete(0, tk.END)
    entry.insert(0, filepath)

def add_watermark():
    filepath = entry.get()
    watermark_text = watermark_entry.get()

    if not filepath or not watermark_text:
        status_label.config(text="Please select a file and enter watermark text.")
        return

    output_file = os.path.splitext(filepath)[0] + "_watermarked.docx"

    doc = Document(filepath)

    for section in doc.sections:
        header = section.header
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        run = paragraph.add_run()
        run.text = watermark_text
        run.font.size = docx.shared.Pt(12)
        run.font.color.rgb = docx.shared.RGBColor(128, 128, 128)

    doc.save(output_file)
    status_label.config(text=f"Watermark added. Saved as {output_file}")

root = tk.Tk()
root.title("Watermark App")

frame = tk.Frame(root)
frame.pack(padx=20, pady=20)

select_button = tk.Button(frame, text="Select File", command=select_file)
select_button.grid(row=0, column=0, pady=5)

entry = tk.Entry(frame, width=50)
entry.grid(row=0, column=1, padx=5, pady=5)

watermark_label = tk.Label(frame, text="Enter Watermark Text:")
watermark_label.grid(row=1, column=0)

watermark_entry = tk.Entry(frame, width=50)
watermark_entry.grid(row=1, column=1, padx=5, pady=5)

add_button = tk.Button(frame, text="Add Watermark", command=add_watermark)
add_button.grid(row=2, columnspan=2, pady=10)

status_label = tk.Label(frame, text="", fg="green")
status_label.grid(row=3, columnspan=2)

root.mainloop()
